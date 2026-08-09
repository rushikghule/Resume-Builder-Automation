"""
Export Agent: renders a StructuredResume into a downloadable DOCX file
using python-docx directly (no external template engine needed for this
simple, clean layout).
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from app.graph.state import GraphState

OUTPUT_DIR = Path(__file__).resolve().parents[3] / "exports"
OUTPUT_DIR.mkdir(exist_ok=True)


def _render_docx(resume: dict, out_path: Path) -> None:
    doc = Document()

    # Name as title
    title = doc.add_heading(resume.get("name", "Your Name"), level=0)

    # Contact line
    contact_parts = [p for p in [resume.get("email"), resume.get("phone")] if p]
    if contact_parts:
        contact = doc.add_paragraph(" | ".join(contact_parts))
        contact.runs[0].font.size = Pt(10)

    # Summary
    if resume.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume["summary"])

    # Skills
    if resume.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(resume["skills"]))

    # Experience
    if resume.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in resume["experience"]:
            header = f"{exp.get('title', '')} — {exp.get('company', '')}"
            dates = f"{exp.get('start_date', '')} – {exp.get('end_date', '')}"
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            p.add_run(f"   ({dates})").italic = True
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # Education
    if resume.get("education"):
        doc.add_heading("Education", level=1)
        for edu in resume["education"]:
            line = f"{edu.get('degree', '')} — {edu.get('institution', '')} ({edu.get('year', '')})"
            doc.add_paragraph(line)

    doc.save(str(out_path))


def export_node(state: GraphState) -> GraphState:
    resume = state.get("tailored_resume") or state.get("structured_resume", {})
    session_id = state.get("session_id", "resume")
    out_path = OUTPUT_DIR / f"{session_id}_tailored.docx"

    try:
        _render_docx(resume, out_path)
        export_path = str(out_path)
    except Exception as e:
        return {**state, "error": f"Export failed: {e}", "stage": "error"}

    return {**state, "export_path": export_path, "stage": "exported"}
